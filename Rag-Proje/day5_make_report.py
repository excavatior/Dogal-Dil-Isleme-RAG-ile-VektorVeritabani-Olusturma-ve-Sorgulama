#!/usr/bin/env python3
# day5_make_report.py — evaluation_report.csv'den Word raporu
import pandas as pd
from docx import Document
from datetime import datetime
import textwrap, os

CSV = "evaluation_report.csv"
OUT = "final_RAG_report.docx"

def main():
    if not os.path.exists(CSV):
        raise SystemExit(f"{CSV} bulunamadı. Önce day5_evaluate.py çalıştır.")

    df = pd.read_csv(CSV)

    doc = Document()
    doc.add_heading("RAG Yöntemi ile Vektör Veritabanı Oluşturma ve Sorgulama", level=1)
    doc.add_paragraph(f"Tarih: {datetime.now():%Y-%m-%d %H:%M}")
    doc.add_paragraph("Kapsam: Metinlerin embedding'e dönüştürülerek Chroma vektör veritabanına aktarılması ve sorgu ile en yakın sonuçların getirilmesi. Referans: Chroma Resmî Dokümantasyon, RAG Yapısı – Pinecone.")

    doc.add_heading("Mimari Özeti", level=2)
    doc.add_paragraph("Data → Embedding (paraphrase-multilingual-MiniLM-L12-v2) → Chroma (rag_docs) → Retrieval → Reranker (cross-encoder/ms-marco-MiniLM-L-6-v2) → LLM (OpenAI/HF).")

    doc.add_heading("Değerlendirme Sonuçları (Özet)", level=2)
    doc.add_paragraph(f"- Soru sayısı: {len(df)}")
    if "latency_ms" in df:
        try:
            mean_ms = int(df["latency_ms"].mean())
            doc.add_paragraph(f"- Ortalama yanıt süresi: ~{mean_ms} ms")
        except Exception:
            pass
    doc.add_paragraph("- Cevaplar sonunda kullanılan kaynaklar [Kaynak: path, Parça: i] biçiminde listelendi.")

    doc.add_heading("Soru–Cevap Tablosu", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Soru", "Cevap (kısaltılmış)", "Süre (ms)", "Kaynaklar"

    for _, r in df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r.get("question",""))
        row[1].text = textwrap.shorten(str(r.get("answer","")), width=300, placeholder=" ...")
        row[2].text = str(r.get("latency_ms",""))
        row[3].text = str(r.get("sources",""))

    doc.add_heading("Notlar & Geliştirme Önerileri", level=2)
    for n in [
        "Chunking 800/100 ile başlandı; 600–1000 aralığı veri yapısına göre denenebilir.",
        "Aday sayısı (candidates) 10–20, top_k 3–5 önerilir.",
        "Kalite için veri kapsamı artırılabilir; domain-özel prompt şablonu eklenebilir.",
        "Performans için caching ve daha hafif LLM’ler (veya OpenAI küçük model) tercih edilebilir."
    ]:
        doc.add_paragraph(f"- {n}")

    doc.save(OUT)
    print(f"[OK] Word rapor oluşturuldu -> {OUT}")

if __name__ == "__main__":
    main()
